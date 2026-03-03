import json
import argparse
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TRUTH_PATH = ROOT / "engine_step1_truth" / "truth_profile.json"

def load_truth():
    return json.loads(TRUTH_PATH.read_text(encoding="utf-8"))

def load_jd(job_id: int) -> str:
    p = ROOT / "engine_step2_tailor" / "outputs" / "jobs" / f"job_{job_id}" / "job_description.txt"
    if not p.exists():
        raise SystemExit(f"Missing JD for job {job_id}. Run: python tools\\attach_jd.py --job {job_id}")
    return p.read_text(encoding="utf-8")

def pick_highlights(truth: dict) -> list[str]:
    # truth-locked: only uses existing bullets, no invention
    bullets = []
    for exp in truth.get("experience", []):
        bullets += exp.get("achievements", [])
    for proj in truth.get("projects", []):
        for d in proj.get("details", []):
            bullets.append(f'{proj.get("name")}: {d}')
    return bullets[:6]

def generate(job_id: int, company: str, role: str):
    truth = load_truth()
    _jd = load_jd(job_id)  # stored for traceability, not used to invent claims

    doc = Document()
    doc.add_paragraph(truth["name"])
    doc.add_paragraph(truth["location"])
    doc.add_paragraph(truth["email"])
    doc.add_paragraph(truth["phone"])
    doc.add_paragraph("")

    doc.add_paragraph(f"Dear Hiring Manager,")
    doc.add_paragraph("")

    doc.add_paragraph(
        f"I am writing to express my interest in the {role} role at {company}. "
        f"I am currently pursuing an M.S. in Cybersecurity (AI Emphasis) at Webster University, expected May 2026."
    )
    doc.add_paragraph("")

    doc.add_paragraph("Relevant highlights from my experience and projects include:")
    for b in pick_highlights(truth):
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "I would welcome the opportunity to discuss how I can contribute to your security operations and support secure development practices."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(truth["name"])

    out_dir = ROOT / "engine_step2_tailor" / "outputs" / "jobs" / f"job_{job_id}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "cover_letter.docx"
    doc.save(out_path)
    print("OK:", out_path)

if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--job", type=int, required=True)
    p.add_argument("--company", required=True)
    p.add_argument("--role", required=True)
    a = p.parse_args()
    generate(a.job, a.company, a.role)