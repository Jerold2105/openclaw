import json
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TRUTH_PATH = ROOT / "engine_step1_truth" / "truth_profile.json"
OUTPUT_DIR = ROOT / "engine_step2_tailor" / "outputs"

def load_truth():
    with open(TRUTH_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def generate_resume():
    data = load_truth()
    doc = Document()

    doc.add_heading(data["name"], level=1)
    doc.add_paragraph(f'{data["location"]} | {data["email"]} | {data["phone"]}')
    doc.add_paragraph(data["linkedin"])
    doc.add_paragraph(data["github"])

    doc.add_heading("Education", level=2)
    for edu in data["education"]:
        doc.add_paragraph(
            f'{edu["degree"]}, {edu["institution"]}'
        )

    doc.add_heading("Experience", level=2)
    for exp in data["experience"]:
        doc.add_paragraph(f'{exp["role"]} - {exp["company"]}')
        for a in exp["achievements"]:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Projects", level=2)
    for proj in data["projects"]:
        doc.add_paragraph(proj["name"])
        for d in proj["details"]:
            doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("Skills", level=2)
    for category, skills in data["skills"].items():
        doc.add_paragraph(f'{category}: {", ".join(skills)}')

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / "generated_resume.docx"
    doc.save(output_path)

    print("Resume generated:", output_path)

if __name__ == "__main__":
    generate_resume()